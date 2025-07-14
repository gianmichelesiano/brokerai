"""
File processing service for extracting text from PDF and DOCX files
"""

import io
import os
import tempfile
import logging
from typing import Tuple, Optional, Dict, Any
import PyPDF2
from docx import Document
from app.config.settings import settings
from app.utils.exceptions import FileProcessingError, raise_file_processing_error

logger = logging.getLogger(__name__)


class FileProcessorService:
    """Service for processing uploaded files and extracting text"""
    
    def __init__(self):
        self.max_file_size = settings.MAX_FILE_SIZE
        self.allowed_extensions = settings.ALLOWED_FILE_TYPES
    
    def validate_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Validate uploaded file before processing
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Dict with validation results
        """
        try:
            file_size = len(file_content)
            file_extension = self._get_file_extension(filename)
            
            errors = []
            warnings = []
            
            # Check file size
            if file_size > self.max_file_size:
                errors.append(
                    f"File troppo grande: {file_size / 1024 / 1024:.1f}MB "
                    f"(massimo {self.max_file_size / 1024 / 1024:.0f}MB)"
                )
            
            # Check file extension
            if file_extension not in self.allowed_extensions:
                errors.append(
                    f"Estensione '{file_extension}' non supportata. "
                    f"Formati accettati: {', '.join(self.allowed_extensions)}"
                )
            
            # Check if file is empty
            if file_size == 0:
                errors.append("Il file è vuoto")
            
            # Test text extraction
            text_extractable = False
            text_preview = None
            estimated_pages = None
            
            if len(errors) == 0:
                try:
                    text, success = self.extract_text_from_file(file_content, filename)
                    text_extractable = success
                    if success and text:
                        text_preview = text[:500] + "..." if len(text) > 500 else text
                        # Estimate pages for PDF
                        if file_extension == "pdf":
                            estimated_pages = self._estimate_pdf_pages(file_content)
                except Exception as e:
                    warnings.append(f"Impossibile estrarre il testo: {str(e)}")
            
            return {
                "valid": len(errors) == 0,
                "errors": errors,
                "warnings": warnings,
                "file_size": file_size,
                "file_extension": file_extension,
                "text_extractable": text_extractable,
                "text_preview": text_preview,
                "estimated_pages": estimated_pages
            }
            
        except Exception as e:
            logger.error(f"Error validating file {filename}: {e}")
            return {
                "valid": False,
                "errors": [f"Errore durante la validazione: {str(e)}"],
                "warnings": [],
                "file_size": len(file_content),
                "file_extension": self._get_file_extension(filename),
                "text_extractable": False,
                "text_preview": None,
                "estimated_pages": None
            }
    
    def extract_text_from_file(self, file_content: bytes, filename: str) -> Tuple[str, bool]:
        """
        Extract text from uploaded file
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (extracted_text, success_flag)
        """
        try:
            file_extension = self._get_file_extension(filename)
            
            if file_extension == "pdf":
                return self._extract_from_pdf(file_content)
            elif file_extension == "docx":
                return self._extract_from_docx(file_content)
            elif file_extension == "doc":
                return "Formato DOC non supportato. Convertire in PDF o DOCX.", False
            else:
                return f"Formato '{file_extension}' non supportato.", False
                
        except Exception as e:
            logger.error(f"Error extracting text from {filename}: {e}")
            return f"Errore nell'estrazione del testo: {str(e)}", False
    
    def get_detailed_extraction_result(
        self, 
        file_content: bytes, 
        filename: str
    ) -> Dict[str, Any]:
        """
        Get detailed text extraction result with metadata
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Dict with detailed extraction results
        """
        import time
        
        start_time = time.time()
        
        try:
            text, success = self.extract_text_from_file(file_content, filename)
            extraction_time = time.time() - start_time
            
            result = {
                "success": success,
                "text": text if success else None,
                "text_length": len(text) if success else 0,
                "extraction_time": extraction_time,
                "errors": None if success else [text],
                "warnings": []
            }
            
            # Add additional metadata for PDF files
            if success and self._get_file_extension(filename) == "pdf":
                try:
                    pages_processed = self._count_pdf_pages(file_content)
                    result["pages_processed"] = pages_processed
                except Exception as e:
                    result["warnings"].append(f"Impossibile contare le pagine: {str(e)}")
            
            return result
            
        except Exception as e:
            extraction_time = time.time() - start_time
            logger.error(f"Error in detailed extraction for {filename}: {e}")
            
            return {
                "success": False,
                "text": None,
                "text_length": 0,
                "extraction_time": extraction_time,
                "errors": [f"Errore nell'estrazione: {str(e)}"],
                "warnings": []
            }
    
    def _extract_from_pdf(self, file_content: bytes) -> Tuple[str, bool]:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Check if PDF is encrypted
            if pdf_reader.is_encrypted:
                return "Il PDF è protetto da password e non può essere elaborato.", False
            
            text_parts = []
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num + 1}: {e}")
                    continue
            
            if not text_parts:
                return "Nessun testo trovato nel PDF. Il file potrebbe contenere solo immagini.", False
            
            full_text = "\n\n".join(text_parts)
            
            # Clean up the text
            full_text = self._clean_extracted_text(full_text)
            
            if len(full_text.strip()) < 50:
                return "Il testo estratto è troppo breve. Il PDF potrebbe contenere principalmente immagini.", False
            
            return full_text, True
            
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {e}")
            return f"Errore nell'estrazione del testo dal PDF: {str(e)}", False
    
    def _extract_from_docx(self, file_content: bytes) -> Tuple[str, bool]:
        """Extract text from DOCX file"""
        try:
            # Create temporary file for docx processing
            with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
                temp_file.write(file_content)
                temp_file_path = temp_file.name
            
            try:
                doc = Document(temp_file_path)
                text_parts = []
                
                # Extract text from paragraphs
                for paragraph in doc.paragraphs:
                    if paragraph.text.strip():
                        text_parts.append(paragraph.text)
                
                # Extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text.strip():
                                row_text.append(cell.text.strip())
                        if row_text:
                            text_parts.append(" | ".join(row_text))
                
                if not text_parts:
                    return "Nessun testo trovato nel documento DOCX.", False
                
                full_text = "\n\n".join(text_parts)
                
                # Clean up the text
                full_text = self._clean_extracted_text(full_text)
                
                if len(full_text.strip()) < 10:
                    return "Il testo estratto è troppo breve.", False
                
                return full_text, True
                
            finally:
                # Clean up temporary file
                try:
                    os.unlink(temp_file_path)
                except Exception:
                    pass
            
        except Exception as e:
            logger.error(f"Error extracting text from DOCX: {e}")
            return f"Errore nell'estrazione del testo dal DOCX: {str(e)}", False
    
    def _clean_extracted_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = []
        for line in text.split('\n'):
            cleaned_line = ' '.join(line.split())
            if cleaned_line:
                lines.append(cleaned_line)
        
        # Join lines with proper spacing
        cleaned_text = '\n'.join(lines)
        
        # Remove excessive newlines
        while '\n\n\n' in cleaned_text:
            cleaned_text = cleaned_text.replace('\n\n\n', '\n\n')
        
        return cleaned_text.strip()
    
    def _get_file_extension(self, filename: str) -> str:
        """Get file extension in lowercase"""
        return filename.lower().split('.')[-1] if '.' in filename else ""
    
    def _estimate_pdf_pages(self, file_content: bytes) -> Optional[int]:
        """Estimate number of pages in PDF"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            return len(pdf_reader.pages)
        except Exception:
            return None
    
    def _count_pdf_pages(self, file_content: bytes) -> int:
        """Count actual pages in PDF"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            return len(pdf_reader.pages)
        except Exception as e:
            raise FileProcessingError(f"Impossibile contare le pagine del PDF: {str(e)}")
    
    def get_file_info(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Get comprehensive file information"""
        try:
            file_size = len(file_content)
            file_extension = self._get_file_extension(filename)
            
            info = {
                "filename": filename,
                "file_size": file_size,
                "file_size_mb": round(file_size / 1024 / 1024, 2),
                "file_extension": file_extension,
                "file_type": self._get_file_type_description(file_extension),
                "is_supported": file_extension in self.allowed_extensions,
                "is_valid_size": file_size <= self.max_file_size
            }
            
            # Add PDF-specific info
            if file_extension == "pdf":
                try:
                    pages = self._count_pdf_pages(file_content)
                    info["pages"] = pages
                    info["estimated_size_per_page"] = round(file_size / pages / 1024, 2) if pages > 0 else 0
                except Exception:
                    info["pages"] = None
                    info["estimated_size_per_page"] = None
            
            return info
            
        except Exception as e:
            logger.error(f"Error getting file info for {filename}: {e}")
            raise FileProcessingError(f"Errore nell'analisi del file: {str(e)}")
    
    def _get_file_type_description(self, extension: str) -> str:
        """Get human-readable file type description"""
        type_map = {
            "pdf": "Documento PDF",
            "docx": "Documento Word (DOCX)",
            "doc": "Documento Word (DOC) - Non supportato"
        }
        return type_map.get(extension, f"File .{extension}")


# Global instance
file_processor = FileProcessorService()
